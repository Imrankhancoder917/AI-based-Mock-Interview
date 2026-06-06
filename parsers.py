from __future__ import annotations

import re
import os
import json
from dataclasses import dataclass
_SKILL_ALIASES: dict[str, str] = {
    "js": "javascript",
    "javascript": "javascript",
    "golang": "go",
    "go": "go",
    "shell": "shell script",
    "shell script": "shell script",
    "tailwind": "tailwind css",
    "tailwind css": "tailwind css",
    "tailwindcss": "tailwind css",
    "rest": "rest api",
    "rest api": "rest api",
    "restful": "restful api",
    "restful api": "restful api",
    "postgres": "postgresql",
    "postgresql": "postgresql",
    "ms sql": "ms sql server",
    "ms sql server": "ms sql server",
    "expressjs": "express",
    "express": "express",
    "express.js": "express",
    "nextjs": "next.js",
    "next.js": "next.js",
    "vue": "vue.js",
    "vue.js": "vue.js",
    "vuejs": "vue.js",
    "reactjs": "react",
    "react": "react",
    "react.js": "react",
    "node.js": "node.js",
    "nodejs": "node.js",
    "rails": "ruby on rails",
    "ruby on rails": "ruby on rails",
    "dotnet": ".net",
    ".net": ".net",
    "swift ui": "swiftui",
    "swiftui": "swiftui",
    "digital ocean": "digitalocean",
    "digitalocean": "digitalocean",
    "cv": "computer vision",
    "computer vision": "computer vision",
    "gen ai": "generative ai",
    "generative ai": "generative ai",
    "huggingface": "hugging face",
    "hugging face": "hugging face",
    "llama index": "llamaindex",
    "llamaindex": "llamaindex",
    "edge tts": "edge-tts",
    "edge-tts": "edge-tts",
    "sklearn": "scikit-learn",
    "scikit-learn": "scikit-learn",
    "torch": "pytorch",
    "pytorch": "pytorch",
    "cv2": "opencv",
    "opencv": "opencv",
    "pil": "pillow",
    "pillow": "pillow",
    "jupyter": "jupyter notebook",
    "jupyter notebook": "jupyter notebook",
    "spark": "apache spark",
    "apache spark": "apache spark",
    "kafka": "apache kafka",
    "apache kafka": "apache kafka",
    "airflow": "apache airflow",
    "apache airflow": "apache airflow",
    "data structures": "data structures & algorithms",
    "data structures & algorithms": "data structures & algorithms",
    "algorithms": "data structures & algorithms",
    "dsa": "data structures & algorithms",
    "data structures and algorithms": "data structures & algorithms",
    "oop": "object oriented programming",
    "object oriented programming": "object oriented programming",
    "oops": "object oriented programming",
    "dbms": "database management systems",
    "database management systems": "database management systems",
    "os": "operating systems",
    "operating systems": "operating systems",
    "cn": "computer networks",
    "computer networks": "computer networks",
    "networking": "computer networks",
    "computer networking": "computer networks",
    "se": "software engineering",
    "software engineering": "software engineering",
    "solid": "solid principles",
    "solid principles": "solid principles",
    "visual studio code": "vs code",
    "vs code": "vs code",
    "intellij": "intellij idea",
    "intellij idea": "intellij idea",
    "cyber security": "cybersecurity",
    "cybersecurity": "cybersecurity",
    "ssl": "ssl/tls",
    "ssl/tls": "ssl/tls",
    "tls": "ssl/tls",
}

from pathlib import Path
from typing import Any, Iterable

import requests


# ---------------------------------------------------------------------------
# Section Aliases
# ---------------------------------------------------------------------------

RESUME_SECTION_ALIASES = {
    "skills": {"skills", "technical skills", "core skills", "competencies", "proficiencies"},
    "projects": {"projects", "selected projects", "project experience", "portfolio"},
    "certifications": {
        "certification", "certifications", "certificate", "certificates",
        "course certifications", "professional certifications", "online certifications",
        "training certifications", "training & certifications", "training and certifications",
        "training  certifications", "courses", "course completion", "completed courses",
        "achievements", "certifications & achievements", "certifications and achievements",
        "certifications  achievements", "awards and certifications", "awards and certification",
        "licenses & certifications", "licenses and certifications", "licenses  certifications",
        "credentials", "professional credentials", "certified courses", "licenses", "accreditations"
    },
    "education": {"education", "academic background", "academics", "qualifications"},
    "experience": {"experience", "work experience", "professional experience", "employment history", "work history"},
    "internships": {"internships", "internship", "intern", "industrial training"},
}

JD_SECTION_ALIASES = {
    "required_skills": {"requirements", "required skills", "what you need", "qualifications"},
    "responsibilities": {"responsibilities", "what you'll do", "what you will do", "role", "about the role"},
    "technologies": {"tech stack", "technologies", "tools", "preferred technologies", "stack"},
}

# ---------------------------------------------------------------------------
# DETERMINISTIC SKILL CLASSIFICATION DICTIONARIES
# ---------------------------------------------------------------------------
# Rules:
#   - HTML / CSS must NEVER go to Programming Languages
#   - DBMS / Computer Networks must NEVER go to Programming Languages
#   - MySQL must NEVER go to Programming Languages
#   - Match is done by exact lowercase lookup FIRST; if not found → Other Technologies
#   - Categories are checked in priority order (most specific first)

_PROGRAMMING_LANGUAGES = {
    "java", "python", "c", "c++", "c#", "javascript", "js", "typescript",
    "go", "golang", "rust", "kotlin", "php", "ruby", "scala", "swift",
    "perl", "dart", "r", "bash", "shell", "vba", "matlab", "haskell",
    "elixir", "clojure", "fortran", "cobol", "assembly", "lua",
}

_WEB_TECHNOLOGIES = {
    "html", "css", "html5", "css3", "bootstrap", "tailwind", "tailwindcss",
    "sass", "less", "jquery", "ajax", "json", "xml", "rest", "rest api",
    "restful", "restful api", "graphql", "websockets", "web sockets",
    "http", "https", "cors", "jwt", "oauth", "webpack", "vite", "babel",
    "responsive design", "pwa", "service workers",
}

_DATABASES = {
    "mysql", "sql", "postgresql", "postgres", "mongodb", "sqlite",
    "oracle", "redis", "cassandra", "dynamodb", "firebase", "firestore",
    "mariadb", "ms sql", "sql server", "neo4j", "elasticsearch", "mssql",
    "couchdb", "influxdb", "supabase", "cockroachdb", "hbase", "aurora",
}

_FRAMEWORKS = {
    "flask", "django", "fastapi", "spring", "spring boot", "springboot",
    "express", "expressjs", "express.js", "next.js", "nextjs", "nuxt.js",
    "nuxtjs", "vue", "vue.js", "vuejs", "react", "reactjs", "react.js",
    "node.js", "nodejs", "sqlalchemy",
    "angular", "angularjs", "laravel", "rails", "ruby on rails",
    "asp.net", "asp", ".net", "dotnet", "hibernate", "mybatis",
    "nest.js", "nestjs", "fastify", "koa", "hapi", "gin", "fiber",
    "actix", "rocket", "axum", "ktor", "struts", "play", "pyramid",
    "tornado", "aiohttp", "starlette", "sanic",
}

_MOBILE_DEVELOPMENT = {
    "android", "ios", "react native", "flutter", "xamarin", "ionic",
    "swift ui", "swiftui", "jetpack compose", "kotlin multiplatform",
    "cordova", "phonegap",
}

_CLOUD_PLATFORMS = {
    "aws", "gcp", "azure", "google cloud", "heroku", "vercel", "netlify",
    "digital ocean", "digitalocean", "cloudflare", "linode", "vultr",
    "ibm cloud", "oracle cloud", "alibaba cloud",
}

_DEVOPS_TOOLS = {
    "docker", "kubernetes", "jenkins", "ci/cd", "github actions",
    "gitlab ci", "travis ci", "circle ci", "circleci", "ansible",
    "terraform", "helm", "prometheus", "grafana", "elk", "splunk",
    "nagios", "nginx", "apache", "linux", "ubuntu", "centos", "debian",
    "vagrant", "puppet", "chef", "maven", "gradle",
}

_AI_ML_TECHNOLOGIES = {
    "machine learning", "deep learning", "artificial intelligence", "ai",
    "natural language processing", "nlp", "computer vision", "cv",
    "llm", "large language models", "generative ai", "gen ai",
    "reinforcement learning", "neural networks", "transformers",
    "bert", "gpt", "stable diffusion", "hugging face", "huggingface",
    "openai", "groq", "anthropic", "langchain", "langgraph", "crewai",
    "autogen", "llamaindex", "llama index", "pinecone", "weaviate",
    "chroma", "chromadb", "qdrant", "faiss", "vector database",
    "whisper", "edge-tts", "edge tts", "speech recognition",
    "text to speech", "tts", "ocr", "easyocr",
}

_DATA_SCIENCE_LIBRARIES = {
    "pandas", "numpy", "scipy", "matplotlib", "seaborn", "plotly",
    "scikit-learn", "sklearn", "tensorflow", "keras", "pytorch",
    "torch", "xgboost", "lightgbm", "catboost", "statsmodels",
    "nltk", "spacy", "gensim", "opencv", "cv2", "pil", "pillow",
    "jupyter", "jupyter notebook", "anaconda", "hadoop", "spark",
    "apache spark", "kafka", "apache kafka", "airflow", "apache airflow",
    "dask", "ray", "mlflow", "wandb", "weights & biases",
}

_CYBER_SECURITY = {
    "cybersecurity", "cyber security", "penetration testing", "pen testing",
    "ethical hacking", "kali linux", "metasploit", "nmap", "wireshark",
    "burp suite", "owasp", "ssl", "tls", "encryption", "cryptography",
    "firewalls", "ids", "ips", "siem", "soc", "network security",
    "vulnerability assessment", "threat modeling",
}

_CORE_SUBJECTS = {
    "data structures", "algorithms", "dsa", "data structures and algorithms",
    "object oriented programming", "oop", "oops",
    "database management systems", "dbms",
    "operating systems", "os",
    "computer networks", "cn", "networking", "computer networking",
    "computer organization", "computer architecture",
    "software engineering", "se",
    "theory of computation", "toc",
    "compiler design", "compilers",
    "discrete mathematics", "discrete maths",
    "digital electronics", "digital logic",
    "microprocessors", "embedded systems",
    "design patterns", "solid principles", "solid",
    "system design", "low level design", "high level design",
    "multithreading", "concurrency", "parallel computing",
}

_GENERAL_TOOLS = {
    "git", "gitlab", "bitbucket",
    "postman", "insomnia", "swagger", "jira",
    "confluence", "trello", "slack", "notion",
    "vs code", "visual studio code", "visual studio",
    "intellij", "pycharm", "eclipse", "xcode",
    "android studio", "vim", "neovim", "emacs",
    "npm", "yarn", "pnpm", "pip", "poetry", "conda",
    "linux", "ubuntu", "windows", "macos",
    "figma", "adobe xd", "sketch", "canva",
    "excel", "google sheets", "tableau", "power bi",
    "latex", "markdown",
}

# Priority-ordered classification buckets
# (checked in this order; first match wins)
_CLASSIFICATION_BUCKETS = [
    # Most specific first to avoid false positive in broader buckets
    ("Core Subjects",       _CORE_SUBJECTS),
    ("Databases",           _DATABASES),
    ("Programming Languages", _PROGRAMMING_LANGUAGES),
    ("Web Technologies",    _WEB_TECHNOLOGIES),
    ("Frameworks",          _FRAMEWORKS),
    ("Mobile Development",  _MOBILE_DEVELOPMENT),
    ("Cloud Platforms",     _CLOUD_PLATFORMS),
    ("DevOps Tools",        _DEVOPS_TOOLS),
    ("AI/ML Technologies",  _AI_ML_TECHNOLOGIES),
    ("Data Science",        _DATA_SCIENCE_LIBRARIES),
    ("Cyber Security",      _CYBER_SECURITY),
    ("Tools",               _GENERAL_TOOLS),
]

# Canonical display names (preserve correct casing from source)
_CANONICAL_NAMES: dict[str, str] = {
    "python": "Python",
    "java": "Java",
    "c": "C",
    "c++": "C++",
    "c#": "C#",
    "javascript": "JavaScript",
    "js": "JS",
    "typescript": "TypeScript",
    "go": "Go",
    "golang": "Golang",
    "rust": "Rust",
    "kotlin": "Kotlin",
    "php": "PHP",
    "ruby": "Ruby",
    "scala": "Scala",
    "swift": "Swift",
    "perl": "Perl",
    "dart": "Dart",
    "r": "R",
    "bash": "Bash",
    "shell": "Shell",
    "matlab": "MATLAB",
    "html": "HTML",
    "css": "CSS",
    "html5": "HTML5",
    "css3": "CSS3",
    "bootstrap": "Bootstrap",
    "tailwind": "Tailwind",
    "tailwindcss": "Tailwindcss",
    "jquery": "jQuery",
    "json": "JSON",
    "xml": "XML",
    "rest": "REST",
    "rest api": "REST API",
    "restful": "Restful",
    "restful api": "RESTful API",
    "graphql": "GraphQL",
    "websockets": "WebSockets",
    "jwt": "JWT",
    "oauth": "OAuth",
    "mysql": "MySQL",
    "sql": "SQL",
    "postgresql": "PostgreSQL",
    "postgres": "Postgres",
    "mongodb": "MongoDB",
    "sqlite": "SQLite",
    "oracle": "Oracle",
    "redis": "Redis",
    "cassandra": "Cassandra",
    "dynamodb": "DynamoDB",
    "firebase": "Firebase",
    "firestore": "Firestore",
    "mariadb": "MariaDB",
    "ms sql": "Ms Sql",
    "sql server": "SQL Server",
    "neo4j": "Neo4j",
    "elasticsearch": "Elasticsearch",
    "flask": "Flask",
    "django": "Django",
    "fastapi": "FastAPI",
    "spring": "Spring",
    "spring boot": "Spring Boot",
    "express": "Express",
    "expressjs": "ExpressJS",
    "express.js": "Express.JS",
    "next.js": "Next.js",
    "nextjs": "NextJS",
    "nuxt.js": "Nuxt.js",
    "vue": "Vue",
    "vue.js": "Vue.js",
    "vuejs": "VueJS",
    "react": "React",
    "reactjs": "ReactJS",
    "react.js": "React.JS",
    "node.js": "Node.js",
    "nodejs": "NodeJS",
    "sqlalchemy": "SQLAlchemy",
    "angular": "Angular",
    "laravel": "Laravel",
    "rails": "Rails",
    "asp.net": "ASP.NET",
    ".net": ".NET",
    "dotnet": "Dotnet",
    "android": "Android",
    "ios": "iOS",
    "react native": "React Native",
    "flutter": "Flutter",
    "swiftui": "SwiftUI",
    "swift ui": "Swift Ui",
    "jetpack compose": "Jetpack Compose",
    "aws": "AWS",
    "gcp": "GCP",
    "azure": "Azure",
    "google cloud": "Google Cloud",
    "heroku": "Heroku",
    "vercel": "Vercel",
    "netlify": "Netlify",
    "digital ocean": "Digital Ocean",
    "digitalocean": "DigitalOcean",
    "docker": "Docker",
    "kubernetes": "Kubernetes",
    "jenkins": "Jenkins",
    "ci/cd": "CI/CD",
    "github actions": "GitHub Actions",
    "terraform": "Terraform",
    "helm": "Helm",
    "prometheus": "Prometheus",
    "grafana": "Grafana",
    "nginx": "Nginx",
    "apache": "Apache",
    "linux": "Linux",
    "machine learning": "Machine Learning",
    "deep learning": "Deep Learning",
    "artificial intelligence": "Artificial Intelligence",
    "ai": "AI",
    "nlp": "NLP",
    "computer vision": "Computer Vision",
    "cv": "CV",
    "llm": "LLM",
    "generative ai": "Generative AI",
    "gen ai": "Gen Ai",
    "hugging face": "Hugging Face",
    "huggingface": "Huggingface",
    "openai": "OpenAI",
    "groq": "Groq",
    "langchain": "LangChain",
    "langgraph": "LangGraph",
    "crewai": "CrewAI",
    "autogen": "AutoGen",
    "llamaindex": "LlamaIndex",
    "llama index": "Llama Index",
    "pinecone": "Pinecone",
    "weaviate": "Weaviate",
    "chroma": "Chroma",
    "chromadb": "ChromaDB",
    "faiss": "FAISS",
    "whisper": "Whisper",
    "edge-tts": "Edge-TTS",
    "edge tts": "Edge Tts",
    "ocr": "OCR",
    "easyocr": "EasyOCR",
    "pandas": "Pandas",
    "numpy": "NumPy",
    "scipy": "SciPy",
    "matplotlib": "Matplotlib",
    "seaborn": "Seaborn",
    "plotly": "Plotly",
    "scikit-learn": "Scikit-learn",
    "sklearn": "Sklearn",
    "tensorflow": "TensorFlow",
    "keras": "Keras",
    "pytorch": "PyTorch",
    "torch": "Torch",
    "xgboost": "XGBoost",
    "lightgbm": "LightGBM",
    "catboost": "CatBoost",
    "nltk": "NLTK",
    "spacy": "spaCy",
    "opencv": "OpenCV",
    "cv2": "Cv2",
    "pil": "Pil",
    "pillow": "Pillow",
    "jupyter": "Jupyter",
    "jupyter notebook": "Jupyter Notebook",
    "hadoop": "Hadoop",
    "spark": "Spark",
    "apache spark": "Apache Spark",
    "kafka": "Kafka",
    "airflow": "Airflow",
    "data structures": "Data Structures",
    "algorithms": "Algorithms",
    "dsa": "DSA",
    "data structures and algorithms": "Data Structures And Algorithms",
    "object oriented programming": "Object Oriented Programming",
    "oop": "OOP",
    "oops": "OOPs",
    "database management systems": "Database Management Systems",
    "dbms": "DBMS",
    "operating systems": "Operating Systems",
    "os": "OS",
    "computer networks": "Computer Networks",
    "cn": "CN",
    "networking": "Networking",
    "computer networking": "Computer Networking",
    "software engineering": "Software Engineering",
    "se": "SE",
    "system design": "System Design",
    "design patterns": "Design Patterns",
    "solid": "Solid",
    "solid principles": "SOLID Principles",
    "multithreading": "Multithreading",
    "concurrency": "Concurrency",
    "git": "Git",
    "gitlab": "GitLab",
    "bitbucket": "Bitbucket",
    "postman": "Postman",
    "swagger": "Swagger",
    "jira": "Jira",
    "confluence": "Confluence",
    "trello": "Trello",
    "vs code": "VS Code",
    "visual studio code": "Visual Studio Code",
    "intellij": "Intellij",
    "pycharm": "PyCharm",
    "eclipse": "Eclipse",
    "xcode": "Xcode",
    "android studio": "Android Studio",
    "npm": "npm",
    "yarn": "Yarn",
    "pip": "pip",
    "figma": "Figma",
    "tableau": "Tableau",
    "power bi": "Power BI",
    "cybersecurity": "Cybersecurity",
    "cyber security": "Cyber Security",
    "penetration testing": "Penetration Testing",
    "ethical hacking": "Ethical Hacking",
    "kali linux": "Kali Linux",
    "metasploit": "Metasploit",
    "wireshark": "Wireshark",
    "owasp": "OWASP",
    "ssl": "SSL",
    "tls": "TLS",
    "encryption": "Encryption",
    "cryptography": "Cryptography",
}

# Words/tokens that should NEVER be treated as skill names
_JUNK_TOKENS = {
    "code", "github", "repository", "repo", "link", "links", "demo", "source code",
    "source", "and job descriptions", "and", "settlement", "click here",
    "https", "http", "url", "urls", "email", "emails",
    "view project", "live demo", "see more", "read more", "details",
    "dear hiring manager", "hiring manager", "dear", "sincerely",
    "regards", "thank you", "company name", "company address",
    "greetings", "signature",
}

_KNOWN_SKILL_KEYS = set().union(*[bucket for _, bucket in _CLASSIFICATION_BUCKETS])

_JD_FORBIDDEN_SKILL_TOKENS = {
    "job", "description", "job description", "department", "engineering",
    "exemption", "status", "exemption status", "non-exempt", "non exempt",
    "summary", "responsibilities", "requirements", "preferred qualifications",
    "qualifications", "benefits", "this", "the", "and/or", "and or",
    "system/configuration", "configuration", "company policies",
    "equal opportunity", "legal disclaimer", "legal disclaimers",
}

_JD_NOISE_HEADINGS = {
    "job", "description", "job description", "department", "summary",
    "responsibilities", "requirements", "preferred qualifications",
    "preferred qualification", "exemption status", "non exempt", "non-exempt",
    "equal opportunity statement", "equal opportunity statements",
    "legal disclaimer", "legal disclaimers", "company policies",
    "company policy", "benefits", "about us", "about the company",
}

_JD_LEGAL_OR_POLICY_MARKERS = {
    "equal opportunity", "eeo", "reasonable accommodation",
    "background check", "employment eligibility", "privacy policy",
    "company policy", "benefits", "medical insurance", "dental insurance",
    "401(k)", "401k", "disclaimer",
}

_KNOWN_ROLE_PATTERNS = [
    "Machine Learning Engineer", "Cyber Security Analyst",
    "Cybersecurity Analyst", "Full Stack Developer", "Frontend Developer",
    "Front End Developer", "Backend Developer", "Back End Developer",
    "Software Engineer", "Software Developer", "Data Scientist",
    "Data Analyst", "AI Engineer", "DevOps Engineer", "Cloud Engineer",
    "QA Engineer", "Quality Assurance Engineer", "Product Manager",
]


# ---------------------------------------------------------------------------
# Validation helpers
# ---------------------------------------------------------------------------

def is_valid_resume_entity(text: str) -> bool:
    if not text:
        return False
    text_clean = text.strip()
    text_lower = text_clean.lower()
    if len(text_clean) < 2 and text_lower not in {"c", "r"}:
        return False

    reject_exact = {
        "programming", "technologies", "technology", "libraries", "tools",
        "libraries/tools", "tools/libraries", "programming languages",
        "programming language", "languages", "web technologies", "web technology",
        "web", "databases", "database", "db", "core subjects", "core subject",
        "core", "frameworks", "framework", "other", "others", "subjects",
        "subject", "key skills", "technical skills", "skills", "skills/tools",
        "tools/frameworks", "frameworks/tools", "databases/tools", "platforms",
        "operating systems", "languages/technologies", "languages/frameworks",
        "professional skills", "projects", "project", "experience", "education",
        "certifications", "certification", "certificates", "certificate",
        "and job descriptions.", "and job descriptions",
    }

    if text_lower in reject_exact:
        return False
    if text_lower in _JUNK_TOKENS:
        return False

    filler_phrases = [
        "and job descriptions", "job description", "dear hiring manager",
        "hiring manager", "company address", "company name",
    ]
    for fp in filler_phrases:
        if fp in text_lower:
            return False

    # Reject location/personal info tokens
    _LOCATIONS = {
        "india", "usa", "uk", "canada", "germany", "france", "australia",
        "california", "texas", "london", "new york", "san francisco",
        "bengaluru", "mumbai", "delhi", "noida", "hyderabad", "pune", "chennai",
        "kolkata", "ahmedabad", "jaipur", "bhopal", "indore", "surat",
        "bangalore", "united states", "united kingdom",
    }
    if text_lower in _LOCATIONS:
        return False

    # Reject email patterns
    if re.search(r"[a-zA-Z0-9_.+-]+@[a-zA-Z0-9-]+\.[a-zA-Z0-9-.]+", text_clean):
        return False

    # Reject phone numbers
    if re.search(r"\+?\d[\d\s\-().]{7,}", text_clean):
        return False

    # Reject URLs
    if re.search(r"https?://|www\.", text_clean, re.IGNORECASE):
        return False

    # If the text is starting with lowercase and contains spaces and ends with period,
    # it's likely a sentence/description
    if re.match(r"^[a-z].*\.$", text_clean):
        return False

    # Fix 5: Reject strings containing em-dash or en-dash (likely subtitle/description fragments)
    if "\u2014" in text_clean or "\u2013" in text_clean:
        # Allow only if it's a very short known token (e.g., "CI/CD")
        if len(text_clean.split()) > 3:
            return False

    return True


def _canonical(skill_lower: str, original: str) -> str:
    """Return the canonical display name for a skill."""
    return _CANONICAL_NAMES.get(skill_lower, original)


# ---------------------------------------------------------------------------
# DETERMINISTIC SKILL CLASSIFICATION
# ---------------------------------------------------------------------------

def classify_skill(skill: str) -> tuple[str, str]:
    """
    Returns (category, display_name) for a skill.
    category is one of the bucket names or 'Other Technologies'.
    display_name preserves or corrects the casing.
    Never puts HTML/CSS/DBMS/MySQL/Computer Networks into Programming Languages.
    """
    s = skill.strip()
    if not s:
        return ("Other Technologies", s)
    s_lower = s.lower()

    # Check in priority order
    for category, bucket in _CLASSIFICATION_BUCKETS:
        if s_lower in bucket:
            return (category, _canonical(s_lower, s))

    # Not matched → Other Technologies
    return ("Other Technologies", s)


def _normalize_skill_lookup_key(skill: str) -> str:
    return re.sub(r"\s+", " ", str(skill).strip().lower().replace("_", " "))


def _is_known_skill(skill: str) -> bool:
    key = _normalize_skill_lookup_key(skill)
    if key in _JD_FORBIDDEN_SKILL_TOKENS:
        return False
    return key in _KNOWN_SKILL_KEYS


def canonical_known_skill(skill: str) -> str | None:
    if not is_valid_resume_entity(skill):
        return None
    key = _normalize_skill_lookup_key(skill)
    if key in _JD_FORBIDDEN_SKILL_TOKENS:
        return None
        
    if key in _KNOWN_SKILL_KEYS:
        return _canonical(key, skill.strip())
        
    # Return unknown skills as well so they can be matched
    return key.title()


def _categorize_skills(skills: list[str]) -> dict[str, list[str]]:
    """
    Takes a flat list of skill strings, classifies each deterministically,
    and returns a dict of category → [skills].
    Fix 4: Uses alias resolution before dedup so DSA/Data Structures/Algorithms
    collapse into a single canonical entry instead of 3 separate ones.
    """
    display_order = [
        "Programming Languages", "Web Technologies", "Databases", "Frameworks",
        "Tools", "Core Subjects", "AI/ML Technologies", "Data Science",
        "Mobile Development", "Cloud Platforms", "DevOps Tools",
        "Cyber Security", "Other Technologies",
    ]
    categories: dict[str, list[str]] = {k: [] for k in display_order}
    seen: set[str] = set()

    for skill in skills:
        s_str = str(skill).strip()
        if not is_valid_resume_entity(s_str):
            continue
        category, display_name = classify_skill(s_str)
        key = display_name.lower()

        # Fix 4: Resolve alias before dedup check so that
        # DSA, Data Structures, Algorithms all map to same canonical key
        alias_key = _SKILL_ALIASES.get(key, key)
        if alias_key in seen or key in seen:
            continue
        seen.add(key)
        seen.add(alias_key)

        if category not in categories:
            categories[category] = []
        categories[category].append(display_name)

    # Sort entries within each category and remove empty ones
    return {k: sorted(v) for k, v in categories.items() if v}


def cleanup_parsed_skills(skills_obj: Any) -> dict[str, list[str]]:
    """
    Normalize the skills object returned by Groq (may be dict or list)
    into a properly classified dict.
    """
    if isinstance(skills_obj, list):
        return _categorize_skills([str(s) for s in skills_obj])

    if isinstance(skills_obj, dict):
        # Flatten all skills from the dict values and re-classify correctly
        flat: list[str] = []
        for skills_list in skills_obj.values():
            if isinstance(skills_list, list):
                flat.extend(str(s).strip() for s in skills_list)
            elif isinstance(skills_list, str):
                flat.append(skills_list.strip())
        return _categorize_skills(flat)

    return {}


# ---------------------------------------------------------------------------
# Junk filters for projects
# ---------------------------------------------------------------------------

_PROJECT_JUNK_WORDS = {
    "code", "github", "repository", "repo", "link", "demo", "source code",
    "source", "settlement", "click here", "view project", "live demo",
    "see more", "read more", "details", "and", "or", "the", "in",
    "project link", "github link", "github repo", "demo link",
}

_PROJECT_JUNK_PATTERNS = [
    r"^https?://",
    r"^www\.",
    r"github\.com",
    r"^[A-Za-z0-9_.+-]+@",
    r"^\d+$",
]

_MONTHS_PATTERN = r"(?:Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Oct|Nov|Dec|January|February|March|April|June|July|August|September|October|November|December)"
_YEAR_PATTERN = r"\b\d{4}\b"
_RANGE_DELIM = r"\s*(?:-|–|—|to)\s*"
_SINGLE_DATE = rf"(?:{_MONTHS_PATTERN}\s+\d{{4}}|\b\d{{4}}\b)"
_DATE_RANGE = rf"{_SINGLE_DATE}(?:{_RANGE_DELIM}(?:{_SINGLE_DATE}|Present|present))?"

_TRAILING_DATE_RE = re.compile(rf"\s+({_DATE_RANGE})$", re.IGNORECASE)
_IS_DATE_RE = re.compile(rf"^(?:{_DATE_RANGE})$", re.IGNORECASE)


def _is_date_like(text: str) -> bool:
    text_clean = text.strip()
    if not text_clean:
        return False
    if _IS_DATE_RE.match(text_clean):
        return True
    text_lower = text_clean.lower()
    months = {
        "jan", "feb", "mar", "apr", "may", "jun", "jul", "aug", "sep", "oct", "nov", "dec",
        "january", "february", "march", "april", "june", "july", "august", "september",
        "october", "november", "december", "present"
    }
    if text_lower in months:
        return True
    if re.match(r"^[\d\s\-–—/]+$", text_clean):
        return True
    return False


def sanitize_entity(entity: Any) -> str | None:
    """Centralized helper to sanitize and validate entities.
    Rejects: dates, months, years, year ranges, and generic project labels.
    Returns cleaned string or None if rejected.
    """
    if not entity:
        return None
    entity_str = str(entity).strip()
    if not entity_str:
        return None

    # Reject if it is date-like
    if _is_date_like(entity_str):
        return None

    # Check common date patterns (months, years, ranges) explicitly
    date_patterns = [
        r"\b(?:Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Oct|Nov|Dec)[a-z]*\s+\d{4}\b",
        r"\b\d{4}\b",
        r"\b\d{4}\s*[-–]\s*\d{4}\b",
        r"\b\d{4}\s*[-–]\s*Present\b"
    ]
    for pat in date_patterns:
        if re.search(pat, entity_str, re.IGNORECASE):
            return None

    # Reject generic project labels
    generic_labels = {
        "web application",
        "ai web application",
        "full stack application",
        "full stack ai web application",
        "machine learning project",
        "software system",
        "web platform",
        "application",
        "platform",
        "system"
    }
    if entity_str.lower() in generic_labels:
        return None

    return entity_str



def extract_project_date_and_name(raw_name: str) -> tuple[str, str | None]:
    raw_name = raw_name.strip()
    match = _TRAILING_DATE_RE.search(raw_name)
    if match:
        date_str = match.group(1).strip()
        cleaned_name = raw_name[:match.start()].strip()
        cleaned_name = re.sub(r"\s*[-–:|]\s*$", "", cleaned_name).strip()
        return cleaned_name, date_str
    return raw_name, None


def _is_valid_project_name(name: str) -> bool:
    if not name or len(name.strip()) < 3:
        return False
    name_clean = name.strip()
    name_lower = name_clean.lower()

    if name_lower in _PROJECT_JUNK_WORDS:
        return False

    for pat in _PROJECT_JUNK_PATTERNS:
        if re.search(pat, name_clean, re.IGNORECASE):
            return False

    if _is_date_like(name_clean):
        return False

    # Reject project entities that are generic labels / categories
    generic_labels = {
        "web application",
        "ai web application",
        "full stack application",
        "full stack ai web application",
        "machine learning project",
        "software system",
        "web platform",
        "application",
        "platform",
        "system"
    }
    if name_lower in generic_labels:
        return False


    # Reject single word generic tokens
    if len(name_clean.split()) == 1 and name_lower in {
        "code", "github", "demo", "link", "repo", "repository",
        "source", "project", "settlement", "and",
    }:
        return False

    # Reject overly long names (likely a description sentence)
    if len(name_clean.split()) > 9:
        return False

    # Reject if starts with lowercase and has > 4 words (sentence pattern)
    if name_clean[0].islower() and len(name_clean.split()) > 4:
        return False

    if not is_valid_resume_entity(name_clean):
        return False

    return True


# ---------------------------------------------------------------------------
# SKILL_TOKENS for fallback extraction
# ---------------------------------------------------------------------------

SKILL_TOKENS = list(_KNOWN_SKILL_KEYS)


@dataclass
class ParsedDocument:
    extracted_text: str
    sections: dict[str, list[str]]


# ---------------------------------------------------------------------------
# Document parsing utilities
# ---------------------------------------------------------------------------

def _lazy_imports():
    import fitz  # type: ignore
    from docx import Document  # type: ignore
    from PIL import Image  # type: ignore
    return fitz, Document, Image


def _clean_lines(text: str) -> list[str]:
    lines = [re.sub(r"\s+", " ", line).strip() for line in text.splitlines()]
    return [line for line in lines if line]


def _normalize_heading(value: str) -> str:
    return re.sub(r"[^a-z0-9\s]", "", value.lower()).strip()


def _section_match(heading: str, aliases: set[str]) -> bool:
    normalized = _normalize_heading(heading)
    return any(normalized == alias or normalized.startswith(alias + " ") for alias in aliases)


def _extract_section_blocks(lines: list[str], aliases: dict[str, set[str]]) -> dict[str, list[str]]:
    sections = {key: [] for key in aliases}
    active_key = None

    for line in lines:
        if _section_match(line, set().union(*aliases.values())):
            for section_name, section_aliases in aliases.items():
                if _section_match(line, section_aliases):
                    active_key = section_name
                    break
            continue

        if active_key:
            if re.fullmatch(r"[A-Z][A-Z0-9\s/&.-]{2,}", line) and len(line.split()) <= 5:
                active_key = None
                continue
            sections[active_key].append(line)

    return {key: value for key, value in sections.items() if value}


def _section_text(lines: list[str], aliases: set[str]) -> str:
    collecting = False
    collected: list[str] = []

    for line in lines:
        if _section_match(line, aliases):
            collecting = True
            continue
        if collecting and re.fullmatch(r"[A-Z][A-Za-z\s/&-]{2,}", line) and len(line.split()) <= 5:
            break
        if collecting:
            collected.append(line)

    return "\n".join(collected).strip()


def _list_from_text(text: str) -> list[str]:
    parts = re.split(r"[\n•·|;/]", text)
    values = []
    for part in parts:
        cleaned = re.sub(r"^[-*]\s*", "", part).strip()
        if cleaned:
            values.append(cleaned)
    return values


def _has_exact_skill_match(skill: str, text: str, raw_text: str | None = None) -> bool:
    """Check if skill exists in text using strict word/token boundaries."""
    skill_lower = skill.strip().lower()
    text_lower = text.lower()
    
    # Special case for "go" - require case-sensitive Go or Golang
    if skill_lower == "go":
        target_text = raw_text if raw_text is not None else text
        return bool(re.search(r"(?<![A-Za-z0-9])(?:Go|Golang)(?![A-Za-z0-9])", target_text))
        
    # Special case for "c" and "r" - require case-sensitive upper C or R
    if skill_lower in {"c", "r"}:
        target_text = raw_text if raw_text is not None else text
        return bool(re.search(r"(?<![A-Za-z0-9])" + re.escape(skill_lower.upper()) + r"(?![A-Za-z0-9+#.])", target_text))

    pattern = r"(?<![a-z0-9+#])" + re.escape(skill_lower) + r"(?![a-z0-9+#])"
    return bool(re.search(pattern, text_lower))


def _extract_skills_from_text(text: str) -> list[str]:
    lowered = text.lower()
    found = []
    for token in sorted(SKILL_TOKENS, key=len, reverse=True):
        if _has_exact_skill_match(token, lowered, text) and token not in found:
            found.append(token)
    for match in re.findall(r"\b[A-Za-z][A-Za-z0-9+/.-]{1,}\b", text):
        lowered_match = match.lower()
        if len(lowered_match) > 2 and lowered_match not in found and re.search(r"[A-Z]|\.|/|\+", match):
            if not any(_has_exact_skill_match(lowered_match, ext) for ext in found):
                found.append(match)
    return found

def _extract_known_skills_from_text(text: str) -> list[str]:
    lowered = text.lower()
    found = []
    for token in sorted(SKILL_TOKENS, key=len, reverse=True):
        if _has_exact_skill_match(token, lowered, text) and token not in found:
            found.append(token)
    return found


def _strip_jd_noise(text: str) -> str:
    cleaned_lines: list[str] = []
    for raw_line in _clean_lines(text):
        normalized = _normalize_heading(raw_line)
        lower = raw_line.lower()
        if normalized in {_normalize_heading(h) for h in _JD_NOISE_HEADINGS}:
            continue
        if lower in _JD_FORBIDDEN_SKILL_TOKENS:
            continue
        if any(marker in lower for marker in _JD_LEGAL_OR_POLICY_MARKERS):
            continue
        cleaned_lines.append(raw_line)
    return "\n".join(cleaned_lines)


def _extract_validated_skills_from_text(text: str) -> list[str]:
    cleaned_text = _strip_jd_noise(text)
    lowered = cleaned_text.lower()
    found: list[str] = []
    seen: set[str] = set()

    for token in sorted(_KNOWN_SKILL_KEYS, key=len, reverse=True):
        if not _has_exact_skill_match(token, lowered, cleaned_text):
            continue
        display = _canonical(token, token)
        key = display.lower()
        if any(re.search(r"\b" + re.escape(key) + r"\b", existing.lower()) for existing in found):
            continue
        if key not in seen:
            seen.add(key)
            found.append(display)

    return found[:30]


def _extract_experience_level_from_text(text: str) -> str:
    """Extract experience level from JD text.

    Fix 2: Check fresher/intern BEFORE manager so that JDs containing
    "hiring manager" or "report to manager" don't get misclassified.
    Also require 'manager' to appear in a job-title context, not just anywhere.
    """
    lower = text.lower()
    # Check fresher/intern FIRST — most specific, avoids false positives
    if re.search(r"\b(junior|jr\.?|entry level|entry-level|fresher|freshers|intern|interns|trainee)\b", lower):
        return "fresher"
    if re.search(r"\b(lead|principal|staff)\s+(engineer|developer|designer|architect)\b", lower):
        return "lead"
    # Only match 'manager' in job-title context, NOT in "hiring manager" / "report to manager"
    if re.search(r"\b(engineering manager|product manager|project manager|team manager|technical manager)\b", lower):
        return "manager"
    if re.search(r"\b(senior|sr\.?)\s+(engineer|developer|designer|analyst)\b", lower):
        return "senior"
    if re.search(r"\b(senior|sr\.?)\b", lower):
        return "senior"
    if re.search(r"\b(mid|intermediate)\b", lower):
        return "mid"
    if re.search(r"\b[3-9]\+?\s*(?:years|yrs)\b|\b(?:[1-9][0-9])\+?\s*(?:years|yrs)\b", lower):
        return "mid"
    return "any"


def extract_text_from_pdf(file_path: Path) -> str:
    fitz, _, _ = _lazy_imports()
    document = fitz.open(file_path)
    parts = []
    for page in document:
        parts.append(page.get_text("text"))
    return "\n".join(parts)


def extract_text_from_docx(file_path: Path) -> str:
    _, Document, _ = _lazy_imports()
    document = Document(str(file_path))
    return "\n".join(paragraph.text for paragraph in document.paragraphs)


def extract_text_from_image(file_path: Path) -> str:
    _, _, Image = _lazy_imports()
    import easyocr  # type: ignore

    reader = easyocr.Reader(["en"], gpu=False)
    image = Image.open(file_path)
    results = reader.readtext(image, detail=0, paragraph=True)
    return "\n".join(str(item) for item in results)


def parse_document(file_path: Path, doc_type: str) -> ParsedDocument:
    suffix = file_path.suffix.lower()
    if suffix == ".pdf":
        extracted_text = extract_text_from_pdf(file_path)
    elif suffix == ".docx":
        extracted_text = extract_text_from_docx(file_path)
    else:
        extracted_text = extract_text_from_image(file_path)

    # Scrub retrieval metadata globally before any parsing occurs
    extracted_text = re.sub(r"\s*\|\s*H:.*?(?=\n|$)", "", extracted_text)
    extracted_text = re.sub(r"\(\d+\)", "", extracted_text)
    extracted_text = re.sub(r"\[Highlight.*?\]|\(Snippet.*?\)", "", extracted_text, flags=re.IGNORECASE)

    lines = _clean_lines(extracted_text)
    section_aliases = RESUME_SECTION_ALIASES if doc_type == "resume" else JD_SECTION_ALIASES
    sections = _extract_section_blocks(lines, section_aliases)

    return ParsedDocument(extracted_text=extracted_text, sections=sections)


# ---------------------------------------------------------------------------
# Project extraction
# ---------------------------------------------------------------------------

# Known technology tokens for project tech-stack enrichment
_PROJECT_TECH_TOKENS = {
    "python", "java", "c++", "c#", "c", "javascript", "typescript", "golang",
    "go", "ruby", "swift", "kotlin", "rust", "php", "sql", "html", "css",
    "postgresql", "mysql", "sqlite", "mongodb", "redis", "firebase", "firestore",
    "flask", "django", "fastapi", "react", "vue", "angular", "node.js", "nodejs",
    "express", "next.js", "docker", "kubernetes", "git", "aws", "gcp", "azure",
    "bootstrap", "jquery", "spring boot", "spring", "terraform", "nginx",
    "rest api", "graphql", "machine learning", "deep learning", "tensorflow",
    "pytorch", "scikit-learn", "pandas", "numpy", "opencv", "langchain",
    "langgraph", "crewai", "openai", "groq", "whisper", "tailwind",
}


def _normalize_tech_name(tech: str) -> str:
    """Normalize a tech string to its canonical display name."""
    t_lower = tech.strip().lower()
    return _canonical(t_lower, tech.strip())


def parse_fallback_projects(projects_block: list[str]) -> list[dict]:
    """
    Extract project name + tech_stack from raw text lines.
    Only returns entries with a valid project name.
    Rejects: 'Code', 'GitHub', 'Settlement', 'and job descriptions', etc.
    """
    projects = []
    if not projects_block:
        return []

    current_project: dict | None = None

    for line in projects_block:
        line = line.strip()
        if not line or len(line) < 4:
            continue

        # Skip junk lines
        if line.lower() in _PROJECT_JUNK_WORDS:
            continue

        is_bullet = line.startswith(("-", "*", "•", "·", "o "))
        is_description = is_bullet or bool(re.search(
            r"\b(developed|implemented|designed|built|created|helped|worked|using|with|improved|integrated|deployed|automated|optimized|managed|performed|achieved|utilized|leveraged|applied|enabled|reduced|increased)\b",
            line.lower()
        ))

        if is_description:
            # Extract tech names from description lines to enrich current project
            if current_project:
                line_lower = line.lower()
                for tech in _PROJECT_TECH_TOKENS:
                    if _has_exact_skill_match(tech, line_lower, line):
                        tech_name = _normalize_tech_name(tech)
                        if tech_name not in current_project["tech_stack"]:
                            current_project["tech_stack"].append(tech_name)
            continue

        # Potential project title line
        name = line
        tech_stack: list[str] = []

        # Try to extract tech stack from parentheses at end
        match_paren = re.search(r"\(([^)]+)\)$", line)
        if match_paren:
            tech_text = match_paren.group(1)
            name = line[: match_paren.start()].strip()
            tech_stack = [t.strip() for t in re.split(r"[,|/]", tech_text) if t.strip()]
        else:
            # Try splitting by common delimiters
            # Fix 1: Added em-dash (\u2014) to the delimiter list
            for delim in (" - ", " : ", " | ", " \u2013 ", " \u2014 "):
                if delim in line:
                    parts = line.split(delim, 1)
                    name = parts[0].strip()
                    # Split by comma, pipe, slash first; then also split remaining items by spaces
                    raw_tech = [t.strip() for t in re.split(r"[,|/]", parts[1]) if t.strip()]
                    tech_stack = []
                    for t in raw_tech:
                        # If a token has multiple space-separated words, try to split into individual techs
                        words = t.split()
                        if len(words) > 1 and all(w[0].isupper() or w in _PROJECT_TECH_TOKENS for w in words):
                            tech_stack.extend(words)
                        else:
                            tech_stack.append(t)
                    break
            else:
                # Also handle em-dash/en-dash without surrounding spaces
                if "\u2014" in line:
                    parts = line.split("\u2014", 1)
                    name = parts[0].strip()
                    raw_tech = [t.strip() for t in re.split(r"[,|/]", parts[1]) if t.strip()]
                    tech_stack = []
                    for t in raw_tech:
                        words = t.split()
                        if len(words) > 1 and all(w[0].isupper() or w in _PROJECT_TECH_TOKENS for w in words):
                            tech_stack.extend(words)
                        else:
                            tech_stack.append(t)
                elif "\u2013" in line:
                    parts = line.split("\u2013", 1)
                    name = parts[0].strip()
                    raw_tech = [t.strip() for t in re.split(r"[,|/]", parts[1]) if t.strip()]
                    tech_stack = []
                    for t in raw_tech:
                        words = t.split()
                        if len(words) > 1 and all(w[0].isupper() or w in _PROJECT_TECH_TOKENS for w in words):
                            tech_stack.extend(words)
                        else:
                            tech_stack.append(t)

        # Clean the name and extract date
        name, proj_date = extract_project_date_and_name(name)
        name = re.sub(r"^[\d.\-*•·\s]+", "", name).strip()
        name = re.sub(r"^[^\w]+|[^\w]+$", "", name).strip()
        name, proj_date_2 = extract_project_date_and_name(name)
        if proj_date_2:
            proj_date = proj_date_2

        generic_project_labels = {
            "full stack ai web application",
            "web application",
            "ai web application",
            "application",
            "system",
            "platform",
            "software system",
            "machine learning project"
        }
        if name.lower().strip() in generic_project_labels:
            continue

        sanitized_name = sanitize_entity(name)
        if not sanitized_name:
            continue
        name = sanitized_name

        if not _is_valid_project_name(name):
            continue

        # Clean and normalize tech stack
        cleaned_tech: list[str] = []
        for t in tech_stack:
            t_clean = t.strip()
            if is_valid_resume_entity(t_clean) and t_clean.lower() not in _PROJECT_JUNK_WORDS:
                cleaned_tech.append(_normalize_tech_name(t_clean))

        # Enrich from name tokens
        name_lower = name.lower()
        for tech in _PROJECT_TECH_TOKENS:
            if _has_exact_skill_match(tech, name_lower, name):
                tech_name = _normalize_tech_name(tech)
                if tech_name not in cleaned_tech:
                    cleaned_tech.append(tech_name)

        current_project = {
            "name": name,
            "type": "",
            "description": "",
            "tech_stack": cleaned_tech,
            "technologies": cleaned_tech,
            "date": proj_date or ""
        }
        projects.append(current_project)


    return projects[:8]


# ---------------------------------------------------------------------------
# Internship extraction
# ---------------------------------------------------------------------------

def parse_fallback_internships(experience_block: list[str], full_text: str) -> list[dict]:
    internships = []
    current_intern = None

    for line in experience_block:
        line = line.strip()
        if not line or len(line) < 2:
            continue

        is_bullet = line.startswith(("-", "*", "•", "·", "o "))
        
        # Check if line looks like a date range
        date_match = re.search(
            r"\b(Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Oct|Nov|Dec|January|February|March|April|"
            r"June|July|August|September|October|November|December)?\s*\d{4}\s*[-–]\s*"
            r"(Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Oct|Nov|Dec|January|February|March|April|"
            r"June|July|August|September|October|November|December|Present|present)?\s*(\d{4}|Present|present)?\b",
            line,
        )

        if is_bullet:
            continue

        if date_match:
            duration = date_match.group(0)
            if not current_intern:
                current_intern = {"company": "Various", "role": "Intern", "duration": duration}
                internships.append(current_intern)
            else:
                current_intern["duration"] = duration
            continue

        name = re.sub(r"^[\d.\-*•·\s]+", "", line).strip()
        sanitized = sanitize_entity(name)
        if not sanitized:
            continue

        is_role = bool(re.search(r"\b(intern|internship|trainee)\b", sanitized.lower()))
        
        if not current_intern or (current_intern.get("company") != "Various" and current_intern.get("role") != "Intern" and current_intern.get("duration")):
            current_intern = {"company": "Various", "role": "Intern", "duration": ""}
            internships.append(current_intern)

        if is_role:
            current_intern["role"] = sanitized
        else:
            current_intern["company"] = sanitized

    valid_internships = []
    for intern in internships:
        if intern["company"] == "Various" and intern["role"] == "Intern":
            continue
        comp_san = sanitize_entity(intern["company"])
        role_san = sanitize_entity(intern["role"])
        if not comp_san or not role_san:
            continue
        valid_internships.append({
            "company": comp_san,
            "role": role_san,
            "duration": intern.get("duration", "Duration not specified") or "Duration not specified"
        })

    return valid_internships


def parse_fallback_employment(experience_block: list[str]) -> list[dict]:
    employment = []
    current_emp = None

    for line in experience_block:
        line = line.strip()
        if not line or len(line) < 2:
            continue

        is_bullet = line.startswith(("-", "*", "•", "·", "o "))
        
        # Check if line looks like a date range
        date_match = re.search(
            r"\b(Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Oct|Nov|Dec|January|February|March|April|"
            r"June|July|August|September|October|November|December)?\s*\d{4}\s*[-–]\s*"
            r"(Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Oct|Nov|Dec|January|February|March|April|"
            r"June|July|August|September|October|November|December|Present|present)?\s*(\d{4}|Present|present)?\b",
            line,
        )

        if is_bullet:
            continue

        if date_match:
            duration = date_match.group(0)
            if not current_emp:
                current_emp = {"company": "Various", "role": "Employee", "duration": duration}
                employment.append(current_emp)
            else:
                current_emp["duration"] = duration
            continue

        name = re.sub(r"^[\d.\-*•·\s]+", "", line).strip()
        sanitized = sanitize_entity(name)
        if not sanitized:
            continue

        is_role = bool(re.search(r"\b(engineer|developer|analyst|manager|consultant|architect|intern|trainee|employee|lead|designer|specialist|programmer|scientist)\b", sanitized.lower()))
        
        if not current_emp or (current_emp.get("company") != "Various" and current_emp.get("role") != "Employee" and current_emp.get("duration")):
            current_emp = {"company": "Various", "role": "Employee", "duration": ""}
            employment.append(current_emp)

        if is_role:
            current_emp["role"] = sanitized
        else:
            current_emp["company"] = sanitized

    valid_employment = []
    for emp in employment:
        if emp["company"] == "Various" and emp["role"] == "Employee":
            continue
        comp_san = sanitize_entity(emp["company"])
        role_san = sanitize_entity(emp["role"])
        if not comp_san or not role_san:
            continue
        valid_employment.append({
            "company": comp_san,
            "role": role_san,
            "duration": emp.get("duration", "Duration not specified") or "Duration not specified"
        })

    return valid_employment




# ---------------------------------------------------------------------------
# Certificate extraction
# ---------------------------------------------------------------------------

_KNOWN_PROVIDERS = {
    "coursera", "udemy", "nptel", "oracle", "cisco", "aws", "google", 
    "microsoft", "ibm", "infosys springboard", "great learning", 
    "scaler", "simplilearn", "aws academy", "infosys", "springboard"
}

def _is_invalid_cert_line(line: str) -> bool:
    line_lower = line.lower()
    if any(marker in line_lower for marker in [
        "view certificate", "view credential", "credential url", "certificate link",
        "credential link", "http://", "https://", "github.com", "linkedin.com",
        "credential id", "credentialid", "verification link", "show credential"
    ]):
        return True
    return False

def _parse_single_certificate(line: str) -> dict:
    # Extract date
    date = ""
    date_match = re.search(
        r"\b(Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Oct|Nov|Dec|January|February|March|April|"
        r"June|July|August|September|October|November|December)?\s*\d{4}\b",
        line,
        re.IGNORECASE
    )
    if date_match:
        date = date_match.group(0)
        line = line.replace(date, "").strip()
        line = re.sub(r"\(\s*\)|\[\s*\]", "", line).strip()

    provider = ""
    name = line

    # 1. Parentheses
    match_paren = re.search(r"\(([^)]+)\)$", name)
    if match_paren:
        inside = match_paren.group(1).strip()
        if inside.lower() in _KNOWN_PROVIDERS or any(kp in inside.lower() for kp in _KNOWN_PROVIDERS) or len(inside) < 25:
            provider = inside
            name = name[:match_paren.start()].strip()
            
    # 2. Colon
    if not provider and ":" in name:
        parts = name.split(":", 1)
        p_cand = parts[0].strip()
        n_cand = parts[1].strip()
        if p_cand.lower() in _KNOWN_PROVIDERS or len(p_cand) < 25:
            provider = p_cand
            name = n_cand
            
    # 3. Delimiters
    if not provider:
        for delim in (" | ", " - ", " \u2013 ", " \u2014 "):
            if delim in name:
                parts = name.split(delim, 1)
                part1 = parts[0].strip()
                part2 = parts[1].strip()
                if part1.lower() in _KNOWN_PROVIDERS:
                    provider = part1
                    name = part2
                elif part2.lower() in _KNOWN_PROVIDERS:
                    provider = part2
                    name = part1
                else:
                    provider = part2
                    name = part1
                break
                
    # 4. by/from/via
    if not provider:
        for sep in (" by ", " from ", " via "):
            if sep in name.lower():
                idx = name.lower().find(sep)
                provider = name[idx + len(sep):].strip()
                name = name[:idx].strip()
                break
                
    # 5. Fallback check for known provider inside the name
    if not provider:
        for kp in sorted(_KNOWN_PROVIDERS, key=len, reverse=True):
            pattern = r"\b" + re.escape(kp) + r"\b"
            match = re.search(pattern, name, re.IGNORECASE)
            if match:
                provider = match.group(0)
                break

    name = re.sub(r"^[\d.\-*•·\s]+", "", name).strip()
    name = re.sub(r"^[^\w]+|[^\w]+$", "", name).strip()
    provider = re.sub(r"^[^\w]+|[^\w]+$", "", provider).strip()

    provider_display = provider
    if provider.lower() in _KNOWN_PROVIDERS:
        for kp in ["Coursera", "Udemy", "NPTEL", "Oracle", "Cisco", "AWS", "Google", "Microsoft", "IBM", "Infosys Springboard", "Great Learning", "Scaler", "Simplilearn", "AWS Academy"]:
            if kp.lower() == provider.lower():
                provider_display = kp
                break

    return {
        "name": name,
        "provider": provider_display or "Self",
        "date": date or "N/A"
    }

def parse_fallback_certificates(certs_block: list[str]) -> list[dict]:
    if not certs_block:
        return []

    cleaned_lines = []
    for line in certs_block:
        line = line.strip()
        if not line or len(line) < 4:
            continue
        if _is_invalid_cert_line(line):
            continue
        cleaned_lines.append(line)

    certs = []
    i = 0
    while i < len(cleaned_lines):
        line = cleaned_lines[i]
        next_line = cleaned_lines[i + 1] if i + 1 < len(cleaned_lines) else None
        
        # Check for multi-line merge
        merged_provider = ""
        if next_line:
            next_line_lower = next_line.lower()
            is_known_p = next_line_lower in _KNOWN_PROVIDERS or any(kp == next_line_lower for kp in _KNOWN_PROVIDERS)
            is_probable_p = len(next_line) < 25 and not any(w in next_line_lower for w in ["certified", "certificate", "foundation", "foundations", "degree", "diploma", "training", "course"])
            
            if is_known_p or is_probable_p:
                merged_provider = next_line
                parsed = _parse_single_certificate(line)
                parsed["provider"] = merged_provider
                
                # Format provider display name
                if merged_provider.lower() in _KNOWN_PROVIDERS:
                    for kp in ["Coursera", "Udemy", "NPTEL", "Oracle", "Cisco", "AWS", "Google", "Microsoft", "IBM", "Infosys Springboard", "Great Learning", "Scaler", "Simplilearn", "AWS Academy"]:
                        if kp.lower() == merged_provider.lower():
                            parsed["provider"] = kp
                            break
                            
                if parsed["name"] and is_valid_resume_entity(parsed["name"]) and parsed["name"].lower() not in _KNOWN_PROVIDERS:
                    certs.append(parsed)
                i += 2
                continue
                
        # Single line parsing
        parsed = _parse_single_certificate(line)
        if parsed["name"] and is_valid_resume_entity(parsed["name"]) and parsed["name"].lower() not in _KNOWN_PROVIDERS:
            certs.append(parsed)
        i += 1

    return certs[:10]


# ---------------------------------------------------------------------------
# JD field cleaning
# ---------------------------------------------------------------------------

_JD_JUNK_PATTERNS = [
    r"[a-zA-Z0-9_.+-]+@[a-zA-Z0-9-]+\.[a-zA-Z0-9-.]+",  # email
    r"https?://[^\s]+",   # URL
    r"www\.[^\s]+",       # URL
    r"linkedin\.com",     # LinkedIn
    r"github\.com",       # GitHub
]

_JD_JUNK_EXACT = {
    "location", "locations", "address", "addresses", "email", "emails",
    "phone", "phone number", "github", "linkedin", "dear hiring manager",
    "company name", "company address", "greetings", "signature",
    "sincerely", "regards", "thank you",
}

_ADDRESS_WORDS = {
    "street", "road", "ave", "avenue", "drive", "suite", "building",
    "floor", "city", "state", "zip", "country",
    "india", "usa", "uk", "canada", "germany", "france", "australia",
    "california", "texas", "london", "new york", "san francisco",
    "bengaluru", "mumbai", "delhi", "noida", "hyderabad", "pune", "chennai",
    "bangalore", "kolkata", "ahmedabad",
}

_PHONE_RE = re.compile(r"\+?\d{1,4}[-.\s]?\(?\d{1,3}\)?[-.\s]?\d{1,4}[-.\s]?\d{1,4}[-.\s]?\d{1,9}")


def _jd_is_junk(text: str) -> bool:
    t_clean = str(text).strip()
    if not t_clean:
        return True
    t_lower = t_clean.lower()

    if t_lower in _JD_JUNK_EXACT:
        return True

    for pat in _JD_JUNK_PATTERNS:
        if re.search(pat, t_clean, re.IGNORECASE):
            return True

    if len(re.findall(r"\d", t_clean)) >= 7 and _PHONE_RE.search(t_clean):
        return True

    # Greeting patterns
    if re.search(r"(?i)\b(dear|hiring manager|greetings|sincerely|regards|thank you|signature)\b", t_clean):
        return True

    # Address patterns
    parts = [p.strip().lower() for p in re.split(r"[,/]", t_clean)]
    if parts and all(p in _ADDRESS_WORDS for p in parts if p):
        return True

    return False


def _clean_jd_role(role: str) -> str:
    value = re.sub(r"^[\d.\-*•·\s]+", "", str(role).strip())
    value = re.sub(r"\s+", " ", value).strip(" :-|")
    if not value:
        return ""
    if _jd_is_junk(value):
        return ""
    if value.lower() in {"role", "target role", "job title", "title", "n/a", "na", "any"}:
        return ""
    if len(value.split()) > 8 or len(value) > 80:
        return ""
    return value


def _extract_role_from_text(text: str) -> str:
    cleaned_text = _strip_jd_noise(text)
    lines = _clean_lines(cleaned_text)

    title_patterns = [
        r"(?i)^\s*(?:job\s+title|title|position|role)\s*[:\-]\s*(.+)$",
        r"(?i)^\s*(?:we\s+are\s+hiring|hiring\s+for)\s*[:\-]?\s*(.+)$",
    ]
    for line in lines[:20]:
        for pattern in title_patterns:
            match = re.search(pattern, line)
            if match:
                candidate = _clean_jd_role(match.group(1))
                if candidate:
                    return candidate

    lower_text = cleaned_text.lower()
    for role in sorted(_KNOWN_ROLE_PATTERNS, key=len, reverse=True):
        if re.search(r"\b" + re.escape(role.lower()) + r"\b", lower_text):
            return role

    return ""


def clean_jd_fields(jd_data: dict) -> dict:
    summary = jd_data.get("summary", jd_data)

    # Clean required_skills
    skills = summary.get("required_skills") or []
    cleaned_skills: list[str] = []
    seen_skills: set[str] = set()
    for raw_skill in skills:
        skill = str(raw_skill).strip()
        canonical = canonical_known_skill(skill)
        if not canonical:
            continue
        key = canonical.lower()
        if key not in seen_skills:
            seen_skills.add(key)
            cleaned_skills.append(canonical)
    summary["required_skills"] = cleaned_skills

    # Clean responsibilities
    # Fix 3: Also reject resume/cover-letter fragments that contaminate JD responsibilities
    _COVER_LETTER_MARKERS = {
        "i have", "i am", "my project", "i developed", "i built", "i designed",
        "i implemented", "i created", "i worked", "i gained", "i learned",
        "my academic", "my experience", "at university", "at college",
        "my internship", "my resume", "my skills", "through my",
        "i possess", "i bring", "i completed", "i successfully",
        "i was responsible", "i contributed", "i collaborated",
        "dear hiring", "dear sir", "dear madam", "sincerely",
        "thank you for", "looking forward", "please find",
    }
    resps = summary.get("responsibilities") or []
    cleaned_resps = []
    for r in resps:
        r_str = str(r).strip()
        if not r_str:
            continue
        if _jd_is_junk(r_str):
            continue
        if _normalize_heading(r_str) in {_normalize_heading(h) for h in _JD_NOISE_HEADINGS}:
            continue
        if any(marker in r_str.lower() for marker in _JD_LEGAL_OR_POLICY_MARKERS):
            continue
        # Fix 3: Reject lines that look like resume/cover-letter content
        if any(marker in r_str.lower() for marker in _COVER_LETTER_MARKERS):
            continue
        # Reject truncated sentence fragments (start mid-word like "ing at...")
        if re.match(r'^[a-z]{1,4}\s', r_str):
            continue
        cleaned_resps.append(r_str)
    summary["responsibilities"] = cleaned_resps

    # Validate role
    role = _clean_jd_role(str(summary.get("role") or ""))
    if not role:
        role = _extract_role_from_text(str(summary.get("_source_text") or ""))
    summary["role"] = role
    summary.pop("_source_text", None)

    # Normalize experience_level
    exp = str(summary.get("experience_level") or "any").strip().lower()
    valid_exps = {"junior", "mid", "senior", "lead", "manager", "fresher", "any"}
    if exp not in valid_exps:
        if "fresh" in exp or "intern" in exp:
            exp = "fresher"
        elif "senior" in exp:
            exp = "senior"
        elif "lead" in exp:
            exp = "lead"
        elif "manager" in exp:
            exp = "manager"
        else:
            exp = "any"
    summary["experience_level"] = exp

    if "summary" in jd_data:
        jd_data["summary"] = summary
    else:
        jd_data = summary

    return jd_data


# ---------------------------------------------------------------------------
# Technology list extraction (for top-level 'technologies' field)
# ---------------------------------------------------------------------------

# ---------------------------------------------------------------------------
# Global Technology Normalization Pipeline
# ---------------------------------------------------------------------------

_ALL_REGISTRY_SETS = [
    _PROGRAMMING_LANGUAGES,
    _WEB_TECHNOLOGIES,
    _DATABASES,
    _FRAMEWORKS,
    _MOBILE_DEVELOPMENT,
    _CLOUD_PLATFORMS,
    _DEVOPS_TOOLS,
    _AI_ML_TECHNOLOGIES,
    _DATA_SCIENCE_LIBRARIES,
    _CYBER_SECURITY,
    _CORE_SUBJECTS,
    _GENERAL_TOOLS,
]

_MASTER_TECHNOLOGY_REGISTRY = set()
for s in _ALL_REGISTRY_SETS:
    for item in s:
        _MASTER_TECHNOLOGY_REGISTRY.add(item.lower())

# Extra allowed / mapping target keywords
_EXTRA_ALLOWED = {
    "javascript", "react", "node.js", "python", "machine learning",
    "ai", "artificial intelligence", "data structures", "oop",
    "computer networks", "database", "dbms", "dsa", "oops", "cn", "db"
}
_MASTER_TECHNOLOGY_REGISTRY.update(_EXTRA_ALLOWED)

_SKILL_MAPPINGS = {
    "js": ["javascript"],
    "reactjs": ["react"],
    "react.js": ["react"],
    "nodejs": ["node.js"],
    "expressjs": ["express"],
    "express.js": ["express"],
    "vuejs": ["vue.js"],
    "vue.js": ["vue.js"],
    "py": ["python"],
    "ml": ["machine learning"],
    "ai/ml": ["ai", "machine learning"],
    "ai": ["artificial intelligence"],
    "dsa": ["data structures"],
    "oops": ["oop"],
    "oop": ["oop"],
    "cn": ["computer networks"],
    "db": ["database"],
    "dbms": ["database"],
}

# Update standard canonical map with display cases
_CANONICAL_NAMES.update({
    "javascript": "JavaScript",
    "js": "JavaScript",
    "node.js": "Node.js",
    "nodejs": "Node.js",
    "react": "React",
    "reactjs": "React",
    "python": "Python",
    "py": "Python",
    "machine learning": "Machine Learning",
    "ml": "Machine Learning",
    "ai": "AI",
    "artificial intelligence": "Artificial Intelligence",
    "data structures": "Data Structures",
    "dsa": "Data Structures",
    "oop": "OOP",
    "oops": "OOP",
    "computer networks": "Computer Networks",
    "cn": "Computer Networks",
    "database": "Database",
    "db": "Database",
    "dbms": "Database",
})

def clean_tech_string(s: str) -> str:
    s = str(s).strip()
    s = re.sub(r'\s+', ' ', s)
    s = s.strip('"\'.,`()[]{}')
    return s.strip()

def split_skill_string(s: str) -> list[str]:
    s_lower = s.lower()
    if s_lower in _MASTER_TECHNOLOGY_REGISTRY or s_lower in _SKILL_MAPPINGS:
        return [s]
    parts = re.split(r'[,/&]', s)
    return [p.strip() for p in parts if p.strip()]

def get_canonical_name(s_lower: str) -> str:
    if s_lower in _CANONICAL_NAMES:
        return _CANONICAL_NAMES[s_lower]
    if s_lower in {"oop", "dbms", "dsa", "os", "cn", "db", "se", "toc", "api", "rest", "jwt", "ssl", "tls", "sql", "css", "html", "php", "nlp", "cv", "tts", "ocr"}:
        return s_lower.upper()
    return _CANONICAL_NAMES.get(s_lower, s_lower.title())

def resolve_parent_child(skills_set: set[str]) -> set[str]:
    resolved = set(skills_set)
    to_remove = set()
    for s1 in skills_set:
        for s2 in skills_set:
            if s1 != s2:
                pattern = r'\b' + re.escape(s1) + r'\b'
                if re.search(pattern, s2):
                    to_remove.add(s1)
    return resolved - to_remove

def _get_technology_aliases(tech: str) -> set[str]:
    aliases = {tech.lower()}
    
    # Add from _SKILL_ALIASES (reverse lookup)
    for k, v in _SKILL_ALIASES.items():
        if v == tech.lower():
            aliases.add(k.lower())
            
    # Add from _SKILL_MAPPINGS (reverse lookup)
    for k, v in _SKILL_MAPPINGS.items():
        if any(item.lower() == tech.lower() for item in v):
            aliases.add(k.lower())
            
    # Explicit mapping of common variations
    variations = {
        "node.js": {"nodejs"},
        "vue.js": {"vuejs", "vue"},
        "react": {"reactjs", "react.js"},
        ".net": {"dotnet"},
        "c#": {"csharp"},
        "c++": {"cpp"},
        "spring boot": {"springboot"},
        "tailwind css": {"tailwindcss", "tailwind"},
        "github actions": {"githubactions"},
        "rest api": {"restapi", "rest", "restful"},
        "postgresql": {"postgres"},
        "ms sql server": {"mssql", "ms sql"},
    }
    
    if tech.lower() in variations:
        aliases.update(variations[tech.lower()])
        
    return aliases


def _is_technology_in_text(tech: str, text: str) -> bool:
    aliases = _get_technology_aliases(tech)
    for alias in aliases:
        if _has_exact_skill_match(alias, text):
            return True
    return False


def extract_technologies(summary_obj: dict, raw_resume_text: str = "") -> list[str]:
    raw_candidates: set[str] = set()

    skills_obj = summary_obj.get("skills") or {}
    if isinstance(skills_obj, dict):
        for skills_list in skills_obj.values():
            if isinstance(skills_list, list):
                for s in skills_list:
                    if s:
                        raw_candidates.add(str(s))
    elif isinstance(skills_obj, list):
        for s in skills_obj:
            if s:
                raw_candidates.add(str(s))

    projects_list = summary_obj.get("projects") or []
    if isinstance(projects_list, list):
        for p in projects_list:
            if isinstance(p, dict):
                tech_stack = p.get("tech_stack") or []
                if isinstance(tech_stack, list):
                    for t in tech_stack:
                        if t:
                            raw_candidates.add(str(t))

    processed_skills: set[str] = set()
    for candidate in raw_candidates:
        cleaned_cand = clean_tech_string(candidate)
        if not cleaned_cand:
            continue
        
        for split_part in split_skill_string(cleaned_cand):
            part_clean = clean_tech_string(split_part)
            if not part_clean:
                continue
            
            part_lower = part_clean.lower()
            
            if part_lower in _SKILL_MAPPINGS:
                mapped_list = _SKILL_MAPPINGS[part_lower]
                for item in mapped_list:
                    processed_skills.add(item.lower())
            else:
                processed_skills.add(part_lower)

    valid_skills: set[str] = set()
    for s in processed_skills:
        if s in _MASTER_TECHNOLOGY_REGISTRY:
            valid_skills.add(s)

    resolved_skills = resolve_parent_child(valid_skills)
    final_techs = [get_canonical_name(s) for s in resolved_skills]
    
    if raw_resume_text:
        validated_techs = []
        for tech in final_techs:
            if _is_technology_in_text(tech, raw_resume_text):
                validated_techs.append(tech)
        final_techs = validated_techs
        
    return sorted(list(set(final_techs)))


# ---------------------------------------------------------------------------
# Fallback resume result builder
# ---------------------------------------------------------------------------

def _resume_result(parsed_document: ParsedDocument) -> dict:
    text = parsed_document.extracted_text
    lines = _clean_lines(text)
    sections = parsed_document.sections

    skills_block = sections.get("skills") or []
    skills_text = "\n".join(skills_block) if skills_block else _section_text(lines, RESUME_SECTION_ALIASES["skills"])
    extracted_skills_list = _extract_skills_from_text(skills_text or text)
    extracted_skills_list.extend(_extract_known_skills_from_text(text))
    
    projects_block = sections.get("projects") or []
    experience_block = (sections.get("experience") or []) + (sections.get("internships") or [])
    certs_block = sections.get("certifications") or []

    # Also extract skills from project tech stacks
    parsed_projects = parse_fallback_projects(projects_block)
    for p in parsed_projects:
        extracted_skills_list.extend(p.get("tech_stack", []))

    # Deduplicate while preserving order
    extracted_skills_list = list(dict.fromkeys(extracted_skills_list))

    classified_skills = _categorize_skills(extracted_skills_list)

    employment = parse_fallback_employment(experience_block)
    intern_kws = ["internship", "intern", "trainee", "industrial training", "apprentice"]
    internships = []
    experience = []
    companies = []
    
    for emp in employment:
        comp = emp.get("company", "")
        role = emp.get("role", "")
        comp_lower = comp.lower()
        role_lower = role.lower()
        
        if comp and comp != "Various" and comp not in companies:
            companies.append(comp)
            
        if any(kw in role_lower or kw in comp_lower for kw in intern_kws):
            internships.append(emp)
        else:
            experience.append(emp)

    summary: dict = {
        "skills": classified_skills,
        "projects": parsed_projects,
        "internships": internships,
        "experience": experience,
        "companies": companies,
        "certificates": parse_fallback_certificates(certs_block),
        "education": _list_from_text("\n".join(sections.get("education", []))),
        "technologies": [],
    }
    summary["technologies"] = extract_technologies(summary, text)

    return {
        "document_type": "resume",
        "summary": summary,
    }


# ---------------------------------------------------------------------------
# Groq-based resume structurer
# ---------------------------------------------------------------------------

def _structure_resume_with_groq(parsed_document: ParsedDocument) -> dict | None:
    api_key = os.environ.get("GROQ_API_KEY", "")
    model = os.environ.get("SECONDARY_LLM_MODEL", "mixtral-8x7b-32768")
    if not api_key:
        return None

    prompt = (
        "You are a strict resume parsing assistant. Extract and normalize the candidate profile from the provided RAW resume text.\n"
        "Return EXACTLY one JSON object and nothing else with the following keys:\n"
        "  skills: a flat array of skill strings only. Do NOT use category names as skills.\n"
        "  projects: array of objects with 'name' (string, 3-8 words, actual project title), 'type' (string or null, e.g. 'Full Stack AI Web Application', 'Machine Learning Model', 'iOS Mobile app'), 'description' (string or null, brief context), 'tech_stack' (array of strings), and 'date' (string or null, e.g. 'Oct 2025' or '2023 - 2024'). Do NOT append the date to the 'name' field; separate it.\n"
        "    NEVER include 'GitHub', 'Code', 'Repository', 'Link', 'Demo', 'Settlement', or any URL as a project name or tech_stack item.\n"
        "  internships: array of objects with 'company' (string), 'role' (string), 'duration' (string).\n"
        "    Only include entries that contain intern/trainee/internship/industrial training keywords.\n"
        "  experience: array of objects with 'company' (string), 'role' (string), 'duration' (string).\n"
        "    Include non-internship work experience / professional employment.\n"
        "  companies: array of strings containing names of all employers/companies/organizations listed in internships or experience.\n"
        "  certificates: array of objects with 'name' (string), 'provider' (string), 'date' (string).\n"
        "  education: array of strings (degree, institution, years).\n"
        "  technologies: array of strings.\n"
        "CRITICAL RULES:\n"
        "  - Do NOT treat Cities, States, Countries, Emails, Phone Numbers, GitHub URLs, LinkedIn URLs, Addresses as skills.\n"
        "  - Do NOT include 'Dear Hiring Manager', greetings, or signatures.\n"
        "  - Do NOT include 'Code', 'GitHub', 'Settlement', 'and job descriptions' as project names.\n"
        "  - Do NOT treat generic project labels (like 'Web Application', 'AI Web Application', 'Full Stack Application', 'Machine Learning Project', 'Software System', 'Web Platform', 'Application', 'Platform', 'System') as project names. Put the actual name in the 'name' field, and the category/label in the 'type' field.\n"
        "  - Do NOT append date tokens (like 'Oct 2025', '2023 - 2024') to the project 'name'. Strip them and place them in the 'date' field.\n"
        "  - Output only valid JSON. Temperature 0.\n"
        "INPUT_TEXT:\n"
    )

    payload_input = {
        "text": parsed_document.extracted_text,
        "sections": parsed_document.sections,
    }

    full_prompt = prompt + json.dumps(payload_input, ensure_ascii=False)

    endpoints = [
        "https://api.groq.com/openai/v1/chat/completions",
        "https://api.groq.com/v1/chat/completions",
    ]

    for url in endpoints:
        try:
            resp = requests.post(
                url,
                headers={"Authorization": f"Bearer {api_key}", "Content-Type": "application/json"},
                json={
                    "model": model,
                    "messages": [{"role": "user", "content": full_prompt}],
                    "temperature": 0.0,
                    "max_tokens": 1200,
                },
                timeout=15,
            )
            resp.raise_for_status()
            body = resp.json()

            candidates: list[str] = []
            if isinstance(body, dict):
                if "choices" in body and isinstance(body["choices"], list):
                    for ch in body["choices"]:
                        if isinstance(ch, dict):
                            msg = ch.get("message") or {}
                            if isinstance(msg, dict):
                                candidates.append(msg.get("content", ""))
                if not candidates:
                    candidates.append(body.get("output", "") or "")

            if not candidates:
                candidates.append(resp.text)

            for txt in candidates:
                if not txt or not isinstance(txt, str):
                    continue
                s = txt.strip()
                # Strip markdown code fences
                if s.startswith("```"):
                    s = re.sub(r"^```[a-z]*\n?", "", s)
                    s = re.sub(r"\n?```$", "", s)
                    s = s.strip()
                try:
                    obj = json.loads(s)
                    required_keys = {"skills", "projects", "internships", "experience", "companies", "certificates", "education", "technologies"}
                    if not required_keys.issubset(obj.keys()):
                        continue

                    # Re-classify skills deterministically regardless of what Groq said
                    raw_skills = obj.get("skills") or []
                    if isinstance(raw_skills, dict):
                        flat_skills: list[str] = []
                        for v in raw_skills.values():
                            if isinstance(v, list):
                                flat_skills.extend(str(x) for x in v)
                    else:
                        flat_skills = [str(s) for s in (raw_skills if isinstance(raw_skills, list) else [])]
                        
                    # STRICT VALIDATION: Skill must physically exist in the resume text
                    text_lower = parsed_document.extracted_text.lower()
                    flat_skills = [s for s in flat_skills if _has_exact_skill_match(s, text_lower, parsed_document.extracted_text)]

                    obj["skills"] = _categorize_skills(flat_skills)

                    # Clean projects
                    cleaned_projects: list[dict] = []
                    for p in obj.get("projects") or []:
                        if isinstance(p, dict):
                            name = str(p.get("name") or "").strip()
                            name, proj_date = extract_project_date_and_name(name)
                            if not proj_date:
                                proj_date = p.get("date") or None
                            if proj_date:
                                proj_date = str(proj_date).strip()
                            name = re.sub(r"^[\d.\-*•·\s]+", "", name).strip()
                            name = re.sub(r"^[^\w]+|[^\w]+$", "", name).strip()
                            name, proj_date_2 = extract_project_date_and_name(name)
                            if proj_date_2:
                                proj_date = proj_date_2
                            
                            san_name = sanitize_entity(name)
                            if not san_name or not _is_valid_project_name(san_name):
                                continue
                            
                            proj_type = str(p.get("type") or "").strip()
                            description = str(p.get("description") or "").strip()
                            proj_type = sanitize_entity(proj_type) or ""
                            
                            tech = p.get("tech_stack") or p.get("technologies") or []
                            if isinstance(tech, str):
                                tech = [t.strip() for t in re.split(r"[,|/]", tech) if t.strip()]
                            cleaned_tech = [
                                _normalize_tech_name(t)
                                for t in tech
                                if isinstance(t, str)
                                and is_valid_resume_entity(t.strip())
                                and t.strip().lower() not in _PROJECT_JUNK_WORDS
                                and _has_exact_skill_match(t, text_lower, parsed_document.extracted_text)
                            ]
                            cleaned_projects.append({
                                "name": san_name,
                                "type": proj_type,
                                "description": description,
                                "tech_stack": cleaned_tech,
                                "technologies": cleaned_tech,
                                "date": proj_date or ""
                            })
                        elif isinstance(p, str):
                            name, proj_date = extract_project_date_and_name(p)
                            name = re.sub(r"^[\d.\-*•·\s]+", "", name).strip()
                            name = re.sub(r"^[^\w]+|[^\w]+$", "", name).strip()
                            name, proj_date_2 = extract_project_date_and_name(name)
                            if proj_date_2:
                                proj_date = proj_date_2
                            
                            san_name = sanitize_entity(name)
                            if san_name and _is_valid_project_name(san_name):
                                cleaned_projects.append({
                                    "name": san_name,
                                    "type": "",
                                    "description": "",
                                    "tech_stack": [],
                                    "technologies": [],
                                    "date": proj_date or ""
                                })
                    obj["projects"] = cleaned_projects

                    # Clean internships
                    cleaned_internships: list[dict] = []
                    intern_kws = ["internship", "intern", "trainee", "industrial training", "apprentice"]
                    for i in obj.get("internships") or []:
                        if not isinstance(i, dict):
                            continue
                        comp = re.sub(r"^[\d.\-*•·\s]+", "", str(i.get("company") or "")).strip()
                        role = re.sub(r"^[\d.\-*•·\s]+", "", str(i.get("role") or "")).strip()
                        dur = str(i.get("duration") or "").strip()
                        
                        comp_san = sanitize_entity(comp)
                        role_san = sanitize_entity(role)
                        if not comp_san or not role_san:
                            continue
                            
                        if len(comp_san.split()) > 8 or len(role_san.split()) > 8:
                            continue
                        role_lower = role_san.lower()
                        if not any(kw in role_lower or kw in comp_san.lower() for kw in intern_kws):
                            if "trainee" not in role_lower and "intern" not in role_lower:
                                continue
                        cleaned_internships.append({
                            "company": comp_san,
                            "role": role_san,
                            "duration": dur or "Duration not specified",
                        })
                    obj["internships"] = cleaned_internships

                    # Clean experience
                    cleaned_experience: list[dict] = []
                    for e in obj.get("experience") or []:
                        if not isinstance(e, dict):
                            continue
                        comp = re.sub(r"^[\d.\-*•·\s]+", "", str(e.get("company") or "")).strip()
                        role = re.sub(r"^[\d.\-*•·\s]+", "", str(e.get("role") or "")).strip()
                        dur = str(e.get("duration") or "").strip()
                        
                        comp_san = sanitize_entity(comp)
                        role_san = sanitize_entity(role)
                        if not comp_san or not role_san:
                            continue
                            
                        if len(comp_san.split()) > 8 or len(role_san.split()) > 8:
                            continue
                        # Filter out intern roles from work experience just in case
                        role_lower = role_san.lower()
                        comp_lower = comp_san.lower()
                        if any(kw in role_lower or kw in comp_lower for kw in intern_kws):
                            continue
                        cleaned_experience.append({
                            "company": comp_san,
                            "role": role_san,
                            "duration": dur or "Duration not specified",
                        })
                    obj["experience"] = cleaned_experience

                    # Clean and compile companies
                    companies_set = set()
                    for c_name in obj.get("companies") or []:
                        if isinstance(c_name, str):
                            c_clean = re.sub(r"^[\d.\-*•·\s]+", "", c_name).strip()
                            c_san = sanitize_entity(c_clean)
                            if c_san and len(c_san.split()) <= 8:
                                companies_set.add(c_san)
                    # Also collect from internships and experience
                    for i in obj["internships"]:
                        if i["company"] != "Various":
                            i_comp_san = sanitize_entity(i["company"])
                            if i_comp_san:
                                companies_set.add(i_comp_san)
                    for e in obj["experience"]:
                        if e["company"] != "Various":
                            e_comp_san = sanitize_entity(e["company"])
                            if e_comp_san:
                                companies_set.add(e_comp_san)
                    obj["companies"] = sorted(list(companies_set))

                    # Clean certificates
                    cleaned_certs: list[dict] = []
                    for c in obj.get("certificates") or []:
                        if isinstance(c, dict):
                            c_name = str(c.get("name") or "").strip()
                            c_provider = str(c.get("provider") or "").strip()
                            c_date = str(c.get("date") or c.get("issue_date") or "").strip()
                            
                            # If name has separators or provider is missing, parse it
                            if c_name and (not c_provider or c_provider.lower() in {"self", "n/a", ""}):
                                parsed = _parse_single_certificate(c_name)
                                c_name = parsed["name"]
                                c_provider = parsed["provider"]
                                if c_date in {"", "N/A"}:
                                    c_date = parsed["date"]
                            
                            # Validate name
                            c_name = re.sub(r"^[\d.\-*•·\s]+", "", c_name).strip()
                            c_name = re.sub(r"^[^\w]+|[^\w]+$", "", c_name).strip()
                            if not is_valid_resume_entity(c_name) or c_name.lower() in _KNOWN_PROVIDERS:
                                continue
                                
                            cleaned_certs.append({
                                "name": c_name,
                                "provider": c_provider or "Self",
                                "date": c_date or "N/A"
                            })
                        elif isinstance(c, str):
                            if _is_invalid_cert_line(c):
                                continue
                            parsed = _parse_single_certificate(c)
                            if parsed["name"] and is_valid_resume_entity(parsed["name"]) and parsed["name"].lower() not in _KNOWN_PROVIDERS:
                                cleaned_certs.append(parsed)
                    obj["certificates"] = cleaned_certs

                    obj["education"] = obj.get("education") or []
                    obj["technologies"] = extract_technologies(obj, parsed_document.extracted_text)

                    return {
                        "document_type": "resume",
                        "summary": obj,
                    }
                except Exception:
                    continue
        except Exception:
            continue

    return None


# ---------------------------------------------------------------------------
# JD fallback result builder
# ---------------------------------------------------------------------------

def _jd_result(parsed_document: ParsedDocument) -> dict:
    text = parsed_document.extracted_text
    cleaned_text = _strip_jd_noise(text)
    lines = _clean_lines(text)
    sections = parsed_document.sections
    required_skills_text = "\n".join(sections.get("required_skills", [])) or _section_text(
        lines, JD_SECTION_ALIASES["required_skills"]
    )
    responsibilities_text = "\n".join(sections.get("responsibilities", [])) or _section_text(
        lines, JD_SECTION_ALIASES["responsibilities"]
    )

    return {
        "document_type": "jd",
        "summary": {
            "role": _extract_role_from_text(text),
            "required_skills": _extract_validated_skills_from_text(required_skills_text or cleaned_text),
            "experience_level": _extract_experience_level_from_text(cleaned_text),
            "responsibilities": _list_from_text(_strip_jd_noise(responsibilities_text)),
            "_source_text": text,
        },
    }


# ---------------------------------------------------------------------------
# Groq-based JD structurer
# ---------------------------------------------------------------------------

def _structure_jd_with_groq(parsed_document: ParsedDocument) -> dict | None:
    api_key = os.environ.get("GROQ_API_KEY", "")
    if not api_key:
        return None

    model = os.environ.get("SECONDARY_LLM_MODEL", "mixtral-8x7b-32768")

    system_prompt = (
        "You are a strict talent acquisition assistant. Extract structured fields from the RAW job description text.\n"
        "Return EXACTLY one JSON object with keys: role (string), required_skills (array), experience_level (string), responsibilities (array).\n"
        "Normalize experience_level to one of: junior, mid, senior, lead, manager, any.\n"
        "Role must be copied from the JD itself. If no role/title exists, use an empty string. Never invent Backend Developer, Software Engineer, or Full Stack Developer.\n"
        "Skills must be actual technical skills only, such as Java, Python, React, Spring Boot, MySQL, Docker, REST API, AWS, Azure, Kubernetes.\n"
        "CRITICAL: Do NOT include Cities, States, Countries, Emails, Phone Numbers, GitHub URLs, LinkedIn URLs, Addresses, "
        "'Dear Hiring Manager', 'Company Name', 'Company Address', 'Greetings', or 'Signatures' "
        "as skills, roles, or responsibilities. Do NOT include document headers, metadata, legal text, benefits, company policy text, "
        "'Job', 'Description', 'Department', 'Exemption Status', 'Non-Exempt', 'Summary', 'This', 'The', or 'and/or'.\n"
        "Be deterministic (temperature 0). Output only valid JSON.\n"
        "INPUT_TEXT:\n"
    )

    payload_input = {
        "text": parsed_document.extracted_text,
        "sections": parsed_document.sections,
    }
    prompt = system_prompt + json.dumps(payload_input, ensure_ascii=False)

    endpoints = [
        "https://api.groq.com/openai/v1/chat/completions",
        "https://api.groq.com/v1/chat/completions",
    ]

    for url in endpoints:
        try:
            resp = requests.post(
                url,
                headers={"Authorization": f"Bearer {api_key}", "Content-Type": "application/json"},
                json={
                    "model": model,
                    "messages": [{"role": "user", "content": prompt}],
                    "temperature": 0.0,
                    "max_tokens": 600,
                },
                timeout=10,
            )
            resp.raise_for_status()
            body = resp.json()

            candidates: list[str] = []
            if isinstance(body, dict):
                if "choices" in body and isinstance(body["choices"], list):
                    for ch in body["choices"]:
                        if isinstance(ch, dict):
                            msg = ch.get("message") or {}
                            if isinstance(msg, dict):
                                candidates.append(msg.get("content", ""))
                if not candidates:
                    candidates.append(resp.text)

            for txt in candidates:
                if not txt or not isinstance(txt, str):
                    continue
                s = txt.strip()
                if s.startswith("```"):
                    s = re.sub(r"^```[a-z]*\n?", "", s)
                    s = re.sub(r"\n?```$", "", s)
                    s = s.strip()
                try:
                    obj = json.loads(s)
                    if all(k in obj for k in ["role", "required_skills", "experience_level", "responsibilities"]):
                        required_skills = [str(s) for s in (obj.get("required_skills") or [])]
                        text_lower = parsed_document.extracted_text.lower()
                        required_skills = [s for s in required_skills if _has_exact_skill_match(s, text_lower, parsed_document.extracted_text)]
                        
                        return {
                            "document_type": "jd",
                            "summary": {
                                "role": str(obj.get("role") or ""),
                                "required_skills": required_skills,
                                "responsibilities": obj.get("responsibilities") or [],
                                "experience_level": str(obj.get("experience_level") or "any"),
                                "_source_text": parsed_document.extracted_text,
                            },
                        }
                except Exception:
                    continue
        except Exception:
            continue

    return None


# ---------------------------------------------------------------------------
# Public entry point
# ---------------------------------------------------------------------------

def parse_uploaded_file(file_path: Path, doc_type: str) -> dict:
    parsed_document = parse_document(file_path, doc_type)
    
    if doc_type == "resume":
        print("\n" + "="*50)
        print("DEBUG LOGGING: RESUME PARSING PIPELINE")
        print("="*50)
        print(f"1. Raw extracted text length: {len(parsed_document.extracted_text)}")
        
        detected_sections = list(parsed_document.sections.keys())
        print(f"2. Detected section names:\n{json.dumps(detected_sections, indent=2)}")
        
        internships_content = parsed_document.sections.get("internships", [])
        print(f"3. Internship section content (len={len(internships_content)}):")
        for line in internships_content[:5]:
            print(f"   - {line}")
            
        certs_content = parsed_document.sections.get("certifications", [])
        print(f"4. Certification section content (len={len(certs_content)}):")
        for line in certs_content[:5]:
            print(f"   - {line}")
            
        fallback_internships = parse_fallback_internships(
            parsed_document.sections.get("experience", []), 
            parsed_document.extracted_text
        )
        print(f"5. Output of parse_fallback_internships():\n   {json.dumps(fallback_internships, indent=2)}")
        
        fallback_certs = parse_fallback_certificates(certs_content)
        print(f"6. Output of parse_fallback_certificates():\n   {json.dumps(fallback_certs, indent=2)}")
        
        structured = _structure_resume_with_groq(parsed_document)
        print(f"7. Output returned by _structure_resume_with_groq():\n   {json.dumps(structured, indent=2) if structured else 'None'}")
        
        final_result = structured if structured is not None else _resume_result(parsed_document)
        print(f"8. Final merged result before saving:\n   {json.dumps(final_result, indent=2)}")
        print("="*50 + "\n")
        
        return final_result

    structured_jd = None
    try:
        structured_jd = _structure_jd_with_groq(parsed_document)
    except Exception:
        structured_jd = None

    if structured_jd is not None:
        return clean_jd_fields(structured_jd)

    return clean_jd_fields(_jd_result(parsed_document))
